import os
import shutil
import pdfplumber
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r"D:\Tesseract-OCR\tesseract.exe"

from fastapi import  APIRouter, UploadFile, File
from docx import Document
from PyPDF2 import PdfReader
from PIL import Image
from pdf2image import convert_from_path

router = APIRouter()

UPLOAD_FOLDER = 'uploads'


def read_txt(path):

    encodings = ["utf-8", "cp1251", "latin-1"]

    for encoding in encodings:
        try:
            with open(path, "r", encoding=encoding) as f:
                return f.read()
        except:
            continue

    return "Could not read file"

def read_docx(path):
    doc = Document(path)

    full_text = []
    #zwykly tekst
    for para in doc.paragraphs:
        full_text.append(para.text)
    #tabeli
    for table in doc.tables:
        for row in table.rows:

            row_data = []

            for cell in row.cells:
                row_data.append(cell.text.strip())
            full_text.append(" | ".join(row_data))

    return "\n".join(full_text)

def read_pdf(path):

    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:

            extracted = page.extract_text()

            if extracted:
                text = extracted + "\n"

    return text

#OCR

def read_image(path):

    image = Image.open(path)

    image = image.convert("L")

    text = pytesseract.image_to_string(
        image,
        lang="rus+eng+pol",
        config="--psm 6"
    )

    return text

def read_scanned_pdf(path):

    pages = convert_from_path(path, dpi=300)

    text = ""

    for page in pages:

        extracted = pytesseract.image_to_string(
            page,
            lang="rus+eng+pol",
            config="--psm 6"
        )

        text += extracted + "\n"

    return text

@router.post('/upload')
async def upload_file(file: UploadFile = File(...)):

    file_path = os.path.join(UPLOAD_FOLDER, file.filename)

    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    content = ""

    if file.filename.endswith(".txt"):
        content = read_txt(file_path)

    elif file.filename.endswith(".docx"):
        content = read_docx(file_path)

    elif file.filename.endswith(".pdf"):
        normal_text = read_pdf(file_path)
        ocr_text = read_scanned_pdf(file_path)
        content = normal_text + "\n" + ocr_text

    elif file.filename.endswith(".png", ".jpg", ".jpeg"):
        content = read_image(file_path)

    return {
        "filename": file.filename,
        "content": content[:3000]
    }